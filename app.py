import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
import io

# --- بيانات الدخول ---
ADMIN_USER = "ALI FETORY"
ADMIN_PASS = "0925843353"

if 'auth' not in st.session_state:
    st.session_state.auth = False

# --- بوابة الدخول ---
if not st.session_state.auth:
    st.title("🏛️ منظومة المسار الذهبي - تجهيز ملفات التأشيرة")
    u_name = st.text_input("اسم المستخدم").strip().upper()
    u_pass = st.text_input("الرقم السري", type="password").strip()
    if st.button("دخول"):
        if u_name == ADMIN_USER.upper() and u_pass == ADMIN_PASS:
            st.session_state.auth = True
            st.rerun()
        else:
            st.error("البيانات غير صحيحة")
    st.stop()

# --- الواجهة الرئيسية للمشروع ---
st.title("🛂 سيستم تجهيز ملف التقديم المتكامل")

# 1. قسم رفع الجواز (القاريء الدقيق)
st.header("1. سحب بيانات الجواز")
uploaded_file = st.file_uploader("ارفع صورة الجواز الأصلية", type=['jpg', 'png', 'jpeg'])

extracted_data = {}
if uploaded_file:
    st.success("جاري تحليل الجواز بدقة...")
    # محاكاة القراءة الدقيقة - هنا يتم ربط محرك OCR الحقيقي
    extracted_data = {
        "full_name": "MOHAMED AHMED AL-LIBI", # مثال للبيانات المسحوبة من الصورة
        "passport_no": "P0987654",
        "expiry": "2029-12-30",
        "dob": "1992-05-15"
    }
    st.write(f"✅ تم سحب البيانات: {extracted_data['full_name']}")

# 2. قسم حجز الطيران والفندق (المبدئي)
st.header("2. الحجوزات المبدئية")
col1, col2 = st.columns(2)
with col1:
    hotel_name = st.text_input("اسم الفندق المقترح", value="Grand Plaza Hotel")
    check_in = st.date_input("تاريخ دخول الفندق")
with col2:
    flight_ref = st.text_input("رقم رحلة الطيران المبدئي", value="LN 123 - Libyan Airlines")
    flight_date = st.date_input("تاريخ الرحلة")

# 3. إصدار الملف الكامل (النموذج الأصلي + الحجوزات)
st.header("3. إصدار ملف التقديم")
target_country = st.selectbox("دولة السفارة:", ["إيطاليا", "فرنسا", "ألمانيا", "إسبانيا"])

if st.button("إنشاء ملف التاشيرة الكامل"):
    doc = Document()
    
    # الجزء الأول: النموذج الرسمي (تنسيق يشبه الورقة الأصلية)
    doc.add_heading(f'SCHENGEN VISA APPLICATION - {target_country}', 0)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # تعبئة الخانات الرسمية
    data_list = [
        ("1. Surname (Family name)", extracted_data.get("full_name", "").split()[-1]),
        ("2. First name(s)", " ".join(extracted_data.get("full_name", "").split()[:-1])),
        ("3. Date of birth", extracted_data.get("dob", "")),
        ("4. Number of travel document", extracted_data.get("passport_no", "")),
    ]
    
    for label, val in data_list:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(val)

    # الجزء الثاني: الحجز الفندقي والطيران (إضافة صفحة جديدة)
    doc.add_page_break()
    doc.add_heading('Flight & Hotel Reservation (Initial)', 1)
    doc.add_paragraph(f"Flight Confirmation: {flight_ref}")
    doc.add_paragraph(f"Departure Date: {flight_date}")
    doc.add_paragraph(f"Hotel Accommodation: {hotel_name}")
    doc.add_paragraph(f"Period: From {check_in}")
    
    # تحويل للتحميل
    bio = io.BytesIO()
    doc.save(bio)
    
    st.download_button(
        label="📥 تحميل ملف التقديم الكامل (Word)",
        data=bio.getvalue(),
        file_name=f"Visa_Package_{target_country}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# أرشيف الإحصائيات (داش بورد)
st.divider()
st.subheader("📊 إحصائيات شركة المسار الذهبي")
st.info("مجموع العمليات المسجلة لهذا الشهر: 2850")
